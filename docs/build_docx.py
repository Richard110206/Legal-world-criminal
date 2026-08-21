"""将 CRIMINAL_ADAPTATION.md 转换为排版良好的 Word 文档。"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

SRC = r"E:\github\Legal-world-criminal\docs\CRIMINAL_ADAPTATION.md"
DST = r"E:\github\Legal-world-criminal\docs\CRIMINAL_ADAPTATION.docx"


def set_cell_shading(cell, color: str):
    """给表格单元格加背景色."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def build():
    with open(SRC, "r", encoding="utf-8") as f:
        text = f.read()

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ── 样式 ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)
    # 设置中文字体回退
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # Heading styles
    for i in range(1, 4):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "微软雅黑"
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if i == 1:
            h.font.size = Pt(18)
        elif i == 2:
            h.font.size = Pt(14)
        else:
            h.font.size = Pt(12)

    # Code style
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.line_spacing = 1.15
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.left_indent = Cm(0.5)

    # ── 解析 Markdown → Word ──
    lines = text.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []
    in_yaml_block = False
    yaml_lines: list[str] = []

    def flush_code():
        nonlocal code_lines, in_code_block
        if not code_lines:
            return
        # 浅灰背景段落模拟代码块
        for cl in code_lines:
            p = doc.add_paragraph(cl, style="CodeBlock")
            # 给段落加背景
            pPr = p._element.get_or_add_pPr()
            shd = pPr.makeelement(qn("w:shd"), {
                qn("w:fill"): "F5F5F5",
                qn("w:val"): "clear",
            })
            pPr.append(shd)
        code_lines = []
        in_code_block = False

    def flush_yaml():
        nonlocal yaml_lines, in_yaml_block
        if not yaml_lines:
            return
        for cl in yaml_lines:
            p = doc.add_paragraph(cl, style="CodeBlock")
            pPr = p._element.get_or_add_pPr()
            shd = pPr.makeelement(qn("w:shd"), {
                qn("w:fill"): "F0F4F0",
                qn("w:val"): "clear",
            })
            pPr.append(shd)
        yaml_lines = []
        in_yaml_block = False

    def flush_table():
        nonlocal table_rows, in_table
        if len(table_rows) < 2:
            table_rows = []
            in_table = False
            return
        rows_count = len(table_rows)
        cols_count = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=rows_count, cols=cols_count, style="Table Grid")
        tbl.autofit = True
        for ri, row_data in enumerate(table_rows):
            for ci, cell_text in enumerate(row_data):
                if ci >= cols_count:
                    break
                cell = tbl.rows[ri].cells[ci]
                cell.text = cell_text.strip()
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = "微软雅黑"
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                if ri == 0:
                    set_cell_shading(cell, "E8E8E8")
        doc.add_paragraph()  # spacing after table
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]

        # 代码块边界
        if line.strip().startswith("```"):
            flush_code()
            flush_yaml()
            lang = line.strip()[3:].strip()
            if lang in ("yaml", "yml"):
                in_yaml_block = True
            elif lang == "python":
                in_code_block = True
            elif lang == "text":
                in_code_block = True
            else:
                in_code_block = True  # default
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if in_yaml_block:
            yaml_lines.append(line)
            i += 1
            continue

        # 表格检测
        if line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                flush_code()
                flush_yaml()
            # skip separator rows like |---|...|
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # 空行
        if not line.strip():
            flush_code()
            flush_yaml()
            doc.add_paragraph()
            i += 1
            continue

        # 标题
        h1 = re.match(r"^## (.+)$", line)
        h2 = re.match(r"^### (.+)$", line)
        h3 = re.match(r"^#### (.+)$", line)
        if h1:
            flush_code()
            flush_yaml()
            doc.add_heading(h1.group(1), level=1)
            i += 1
            continue
        if h2:
            flush_code()
            flush_yaml()
            doc.add_heading(h2.group(1), level=2)
            i += 1
            continue
        if h3:
            flush_code()
            flush_yaml()
            doc.add_heading(h3.group(1), level=3)
            i += 1
            continue

        # 水平线
        if line.strip() == "---":
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # 普通段落：处理内联代码 `...` 和粗体 **...**
        flush_code()
        flush_yaml()
        p = doc.add_paragraph()
        # 简单内联处理
        remaining = line
        while remaining:
            code_match = re.match(r"^(.*?)`([^`]+)`", remaining)
            bold_match = re.match(r"^(.*?)\*\*(.+?)\*\*", remaining)
            bold2_match = re.match(r"^(.*?)__(.+?)__", remaining)

            if code_match:
                if code_match.group(1):
                    run = p.add_run(code_match.group(1))
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run = p.add_run(code_match.group(2))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                remaining = remaining[code_match.end():]
            elif bold_match:
                if bold_match.group(1):
                    run = p.add_run(bold_match.group(1))
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run = p.add_run(bold_match.group(2))
                run.bold = True
                run.font.name = "微软雅黑"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                remaining = remaining[bold_match.end():]
            elif bold2_match:
                if bold2_match.group(1):
                    run = p.add_run(bold2_match.group(1))
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run = p.add_run(bold2_match.group(2))
                run.bold = True
                run.font.name = "微软雅黑"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                remaining = remaining[bold2_match.end():]
            else:
                run = p.add_run(remaining)
                run.font.name = "微软雅黑"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                remaining = ""

        i += 1

    # 收尾
    flush_code()
    flush_yaml()
    flush_table()

    doc.save(DST)
    print(f"Done → {DST}")


if __name__ == "__main__":
    build()
